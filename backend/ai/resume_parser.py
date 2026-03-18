import PyPDF2
from docx import Document
import os

def extract_resume_text(file_path):
    """
    Extract text from resume files (PDF or DOCX).
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Extracted text from the resume
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file format is not supported
    """
    
    # Check if file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Only PDF and DOCX are supported.")
    except Exception as e:
        raise Exception(f"Error extracting text from {file_path}: {str(e)}")


def extract_text_from_pdf(file_path):
    """
    Extract text from PDF file using PyPDF2.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from PDF
    """
    text = ""
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # Extract text from each page
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")


def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from DOCX
    """
    text = ""
    
    try:
        doc = Document(file_path)
        
        # Extract text from each paragraph
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")


def get_file_info(file_path):
    """
    Get basic information about the resume file.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        dict: File information including name, size, and type
    """
    if not os.path.exists(file_path):
        return None
    
    file_stats = os.stat(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    return {
        'filename': os.path.basename(file_path),
        'size_bytes': file_stats.st_size,
        'size_kb': round(file_stats.st_size / 1024, 2),
        'file_type': file_extension,
        'full_path': file_path
    }
